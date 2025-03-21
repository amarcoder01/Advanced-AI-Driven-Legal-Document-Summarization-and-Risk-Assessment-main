import streamlit as st
import io
import logging
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, BinaryIO

class ExportHandler:
    def __init__(self):
        """Initialize export handler."""
        pass

    def generate_txt(self, content: str) -> str:
        """Generate a well-formatted text document."""
        lines = content.split("\n")
        formatted_lines = []
        
        # Add header
        formatted_lines.append("=" * 80)
        formatted_lines.append("LEGAL DOCUMENT ANALYSIS REPORT")
        formatted_lines.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        formatted_lines.append("=" * 80)
        formatted_lines.append("")
        
        # Process content with proper spacing
        in_section = False
        for line in lines:
            if "DOCUMENT SUMMARY" in line or "RISK ANALYSIS" in line or "RISK SCORE" in line:
                if in_section:
                    formatted_lines.append("")
                
                in_section = True
                section_title = line.strip()
                
                formatted_lines.append("")
                formatted_lines.append(section_title)
                formatted_lines.append("-" * len(section_title))
            elif line.strip() == "=" * 50:
                continue
            else:
                if line.strip():
                    formatted_lines.append(line)
        
        # Add footer
        formatted_lines.append("")
        formatted_lines.append("-" * 80)
        formatted_lines.append("End of Report")
        
        return "\n".join(formatted_lines)

    def generate_docx(self, content: str) -> BinaryIO:
        """Generate a professionally formatted Word document."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Legal Document Analysis"
        doc.core_properties.author = "AI Legal Document Assistant"
        
        # Add header with title
        header = doc.add_heading("Legal Document Analysis Report", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph("_" * 60)
        
        # Process content by sections
        lines = content.split("\n")
        current_section = None
        
        for line in lines:
            if any(section in line for section in [
                "DOCUMENT SUMMARY", "RISK SCORE", "RISK ANALYSIS", 
                "COMPLIANCE ANALYSIS", "DETAILED COMPLIANCE REPORT",
                "KEY FINDINGS", "RECOMMENDATIONS", 
                "CHAT INTERACTION HISTORY", "DOCUMENT COMPARISON ANALYSIS"
            ]):
                doc.add_paragraph()
                heading_text = line.strip().title()
                heading = doc.add_heading(heading_text, level=2)
                current_section = line.strip()
            elif line.strip() and not line.startswith("="):
                p = doc.add_paragraph()
                
                if current_section == "RISK ANALYSIS":
                    if line.lower().startswith(("high", "medium", "low")) and ":" in line:
                        priority, description = line.split(":", 1)
                        priority = priority.strip()
                        
                        priority_run = p.add_run(f"{priority}: ")
                        priority_run.bold = True
                        
                        if "high" in priority.lower():
                            priority_run.font.color.rgb = RGBColor(255, 0, 0)
                        elif "medium" in priority.lower():
                            priority_run.font.color.rgb = RGBColor(255, 165, 0)
                        
                        p.add_run(description.strip())
                    else:
                        p.add_run(line)
                elif current_section == "DOCUMENT COMPARISON ANALYSIS":
                    if "Similarity Score:" in line:
                        p.add_run(line).bold = True
                    elif any(header in line for header in ["Analysis:", "Key Differences:", "Common Elements:"]):
                        p.add_run(line).bold = True
                    else:
                        p.add_run(line)
                elif current_section == "CHAT INTERACTION HISTORY":
                    if ":" in line:  # This is a chat message
                        role, message = line.split(":", 1)
                        role_run = p.add_run(f"{role}: ")
                        role_run.bold = True
                        p.add_run(message.strip())
                    else:
                        p.add_run(line)
                else:
                    p.add_run(line)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph("End of Report")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_pdf(self, content: str) -> BinaryIO:
        """Generate a professionally formatted PDF document."""
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 16)
                self.cell(0, 10, 'Legal Document Analysis Report', 0, 1, 'C')
                self.set_font('Arial', 'I', 10)
                self.cell(0, 5, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, 'C')
                self.ln(5)
                self.line(10, 25, 200, 25)
                self.ln(10)
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')
        
        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        content_parts = content.split("\n\n")
        current_section = None
        
        for part in content_parts:
            if "DOCUMENT SUMMARY" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Document Summary", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                summary_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("DOCUMENT SUMMARY", "").strip()
                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(summary_content))
                pdf.ln(5)
            elif "RISK SCORE" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Risk Assessment Score", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                score_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("RISK SCORE", "").strip()
                for line in score_content.split("\n"):
                    if "Score:" in line or "Risk Level:" in line:
                        pdf.set_font("Arial", "B", 11)
                        pdf.cell(0, 7, line, 0, 1)
                        pdf.set_font("Arial", "", 11)
                    else:
                        pdf.cell(0, 7, line, 0, 1)
                pdf.ln(5)
            elif "RISK ANALYSIS" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Detailed Risk Analysis", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                analysis_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("RISK ANALYSIS", "").strip()
                if "HIGH PRIORITY RISKS" in analysis_content:
                    sections = analysis_content.split("\n\n")
                    for section in sections:
                        if "HIGH PRIORITY RISKS" in section:
                            pdf.set_font("Arial", "B", 12)
                            pdf.set_text_color(255, 0, 0)
                            pdf.cell(0, 10, "High Priority Risks", 0, 1)
                            pdf.set_text_color(0, 0, 0)
                            pdf.set_font("Arial", "", 11)
                            for line in section.split("\n"):
                                if not line.strip() or "HIGH PRIORITY RISKS" in line:
                                    continue
                                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                        elif "MEDIUM PRIORITY RISKS" in section:
                            pdf.set_font("Arial", "B", 12)
                            pdf.set_text_color(255, 165, 0)
                            pdf.cell(0, 10, "Medium Priority Risks", 0, 1)
                            pdf.set_text_color(0, 0, 0)
                            pdf.set_font("Arial", "", 11)
                            for line in section.split("\n"):
                                if not line.strip() or "MEDIUM PRIORITY RISKS" in line:
                                    continue
                                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                        elif "LOW PRIORITY RISKS" in section:
                            pdf.set_font("Arial", "B", 12)
                            pdf.set_text_color(0, 128, 0)
                            pdf.cell(0, 10, "Low Priority Risks", 0, 1)
                            pdf.set_text_color(0, 0, 0)
                            pdf.set_font("Arial", "", 11)
                            for line in section.split("\n"):
                                if not line.strip() or "LOW PRIORITY RISKS" in line:
                                    continue
                                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                else:
                    pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(analysis_content))
                pdf.ln(5)
            elif "DETAILED COMPLIANCE REPORT" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Detailed Compliance Report", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                compliance_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("DETAILED COMPLIANCE REPORT", "").strip()
                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(compliance_content))
                pdf.ln(5)
            elif "CHAT INTERACTION HISTORY" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Chat Interaction History", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                chat_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("CHAT INTERACTION HISTORY", "").strip()
                for line in chat_content.split("\n"):
                    if ":" in line:  # This is a chat message
                        role, message = line.split(":", 1)
                        pdf.set_font("Arial", "B", 11)
                        pdf.cell(40, 7, role + ":", 0, 0)
                        pdf.set_font("Arial", "", 11)
                        pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(message.strip()))
                    else:
                        pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                pdf.ln(5)
            elif "DOCUMENT COMPARISON ANALYSIS" in part:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Document Comparison Analysis", 0, 1, "L")
                pdf.set_font("Arial", "", 11)
                comparison_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("DOCUMENT COMPARISON ANALYSIS", "").strip()
                for line in comparison_content.split("\n"):
                    if any(header in line for header in ["Similarity Score:", "Analysis:", "Key Differences:", "Common Elements:"]):
                        pdf.set_font("Arial", "B", 11)
                        pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                        pdf.set_font("Arial", "", 11)
                    else:
                        pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(line))
                pdf.ln(5)
            else:
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 7, self._sanitize_text_for_pdf(part))
                pdf.ln(3)
        
        # Get PDF as bytes
        pdf_str = pdf.output(dest="S")
        if isinstance(pdf_str, str):
            pdf_bytes = pdf_str.encode("latin1")
        else:
            pdf_bytes = pdf_str
            
        buffer = io.BytesIO(pdf_bytes)
        buffer.seek(0)
        return buffer

    def _sanitize_text_for_pdf(self, text: str) -> str:
        """Clean and sanitize text for PDF output."""
        if not text:
            return ""
        
        replacements = {
            '\u2013': '-',    # en dash
            '\u2014': '--',   # em dash
            '\u2018': "'",    # left single quote
            '\u2019': "'",    # right single quote
            '\u201c': '"',    # left double quote
            '\u201d': '"',    # right double quote
            '\u2022': '*',    # bullet
            '\u2026': '...', # ellipsis
            '\u00a0': ' ',   # non-breaking space
        }
        
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        clean_text = ""
        for char in text:
            if ord(char) < 256:
                clean_text += char
            else:
                clean_text += '?'
                
        return clean_text

    def export_ui(self) -> None:
        """Display export options in the UI."""
        st.markdown("### 📥 Export Analysis")
        
        if not st.session_state.get("extracted_text"):
            st.info("Upload and analyze a document to enable export options.")
            return
            
        # Component selection
        st.markdown("#### Select Components to Export")
        
        available_components = []
        if st.session_state.get("summary"):
            available_components.append("Document Summary")
        if st.session_state.get("risks"):
            available_components.append("Risk Analysis")
        if st.session_state.get("risk_score"):
            available_components.append("Risk Score")
        if st.session_state.get("compliance_results"):
            available_components.append("Compliance Analysis")
        if st.session_state.get("key_findings"):
            available_components.append("Key Findings")
        if st.session_state.get("recommendations"):
            available_components.append("Recommendations")
        if st.session_state.get("chat_history"):
            available_components.append("Chat History")
        if st.session_state.get("comparison_results"):
            available_components.append("Document Comparison")
        if st.session_state.get("compliance_details"):
            available_components.append("Detailed Compliance Report")
            
        if not available_components:
            st.info("Analyze the document first to enable export options.")
            return
            
        selected_components = st.multiselect(
            "Choose sections to include:",
            options=available_components,
            default=available_components,
            help="Select which analysis components you want to include in the export"
        )
        
        if not selected_components:
            st.warning("Please select at least one component to export.")
            return
            
        # Prepare content based on selection
        content = self._prepare_selected_content(selected_components)
        
        if content:
            export_format = st.selectbox(
                "Export Format",
                ["PDF", "DOCX", "TXT"],
                help="Choose the format for your exported analysis"
            )
            
            if st.button("Export Analysis"):
                try:
                    if export_format == "PDF":
                        buf = self.generate_pdf(content)
                        st.download_button(
                            label="Download PDF",
                            data=buf,
                            file_name=f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                            mime="application/pdf"
                        )
                    elif export_format == "DOCX":
                        buf = self.generate_docx(content)
                        st.download_button(
                            label="Download DOCX",
                            data=buf,
                            file_name=f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:  # TXT
                        txt_content = self.generate_txt(content)
                        st.download_button(
                            label="Download TXT",
                            data=txt_content,
                            file_name=f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                            mime="text/plain"
                        )
                except Exception as e:
                    st.error(f"Error generating export: {str(e)}")
                    logging.error(f"Export error: {str(e)}", exc_info=True)

    def _prepare_selected_content(self, selected_components: list) -> Optional[str]:
        """Prepare content based on selected components."""
        content_parts = []
        
        # Add header
        content_parts.extend([
            "LEGAL DOCUMENT ANALYSIS REPORT",
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "=" * 80,
            "\n"
        ])
        
        # Add selected components
        if "Document Summary" in selected_components and st.session_state.get("summary"):
            content_parts.extend([
                "DOCUMENT SUMMARY",
                "=" * 50,
                st.session_state.summary,
                "\n"
            ])
        
        if "Risk Analysis" in selected_components and st.session_state.get("risks"):
            content_parts.extend([
                "RISK ANALYSIS",
                "=" * 50,
                st.session_state.risks,
                "\n"
            ])
        
        if "Risk Score" in selected_components and st.session_state.get("risk_score"):
            content_parts.extend([
                "RISK SCORE ANALYSIS",
                "=" * 50,
                f"Overall Risk Score: {st.session_state.risk_score}/100",
                f"Risk Level: {st.session_state.risk_level}",
                "\n"
            ])
        
        if "Compliance Analysis" in selected_components and st.session_state.get("compliance_results"):
            content_parts.extend([
                "COMPLIANCE ANALYSIS",
                "=" * 50,
                st.session_state.compliance_results,
                "\n"
            ])
        
        if "Detailed Compliance Report" in selected_components and st.session_state.get("compliance_details"):
            content_parts.extend([
                "DETAILED COMPLIANCE REPORT",
                "=" * 50,
                st.session_state.compliance_details,
                "\n"
            ])
        
        if "Key Findings" in selected_components and st.session_state.get("key_findings"):
            content_parts.extend([
                "KEY FINDINGS",
                "=" * 50,
                st.session_state.key_findings,
                "\n"
            ])
        
        if "Recommendations" in selected_components and st.session_state.get("recommendations"):
            content_parts.extend([
                "RECOMMENDATIONS",
                "=" * 50,
                st.session_state.recommendations,
                "\n"
            ])
            
        if "Chat History" in selected_components and st.session_state.get("chat_history"):
            content_parts.extend([
                "CHAT INTERACTION HISTORY",
                "=" * 50,
                self._format_chat_history(st.session_state.chat_history),
                "\n"
            ])
            
        if "Document Comparison" in selected_components and st.session_state.get("comparison_results"):
            content_parts.extend([
                "DOCUMENT COMPARISON ANALYSIS",
                "=" * 50,
                self._format_comparison_results(st.session_state.comparison_results),
                "\n"
            ])
        
        # Add footer
        content_parts.extend([
            "-" * 80,
            "End of Report",
            "\n"
        ])
        
        return "\n".join(content_parts) if content_parts else None

    def _format_chat_history(self, chat_history: list) -> str:
        """Format chat history for export."""
        formatted_chat = []
        for msg in chat_history:
            if isinstance(msg, dict):
                role = msg.get("role", "")
                content = msg.get("content", "")
                formatted_chat.append(f"{role.upper()}: {content}\n")
            elif isinstance(msg, (list, tuple)) and len(msg) >= 2:
                formatted_chat.append(f"{msg[0].upper()}: {msg[1]}\n")
            else:
                formatted_chat.append(str(msg) + "\n")
        return "\n".join(formatted_chat)

    def _format_comparison_results(self, comparison_results: dict) -> str:
        """Format comparison results for export."""
        formatted_comparison = []
        
        if isinstance(comparison_results, dict):
            # Add similarity score if available
            if "similarity" in comparison_results:
                formatted_comparison.append(f"Similarity Score: {comparison_results['similarity']}%\n")
            
            # Add main analysis
            if "analysis" in comparison_results:
                formatted_comparison.append("Analysis:")
                formatted_comparison.append(comparison_results["analysis"])
                formatted_comparison.append("")
            
            # Add key differences
            if "key_differences" in comparison_results:
                formatted_comparison.append("Key Differences:")
                for category, differences in comparison_results["key_differences"].items():
                    formatted_comparison.append(f"\n{category}:")
                    for diff in differences:
                        if isinstance(diff, dict):
                            if "doc1" in diff and "doc2" in diff:
                                formatted_comparison.append("Document 1: " + diff["doc1"])
                                formatted_comparison.append("Document 2: " + diff["doc2"])
                                formatted_comparison.append("-" * 40)
                            elif "content" in diff:
                                formatted_comparison.append(diff["content"])
                        else:
                            formatted_comparison.append(str(diff))
            
            # Add common elements
            if "common_elements" in comparison_results:
                formatted_comparison.append("\nCommon Elements:")
                for category, elements in comparison_results["common_elements"].items():
                    formatted_comparison.append(f"\n{category}:")
                    for element in elements:
                        if isinstance(element, dict):
                            if "title" in element:
                                formatted_comparison.append(f"- {element['title']}")
                            if "content" in element:
                                formatted_comparison.append(f"  {element['content']}")
                        else:
                            formatted_comparison.append(f"- {element}")
        
        return "\n".join(formatted_comparison) 