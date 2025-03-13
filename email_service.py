import smtplib
import streamlit as st
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import io
import logging
from fpdf import FPDF
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(level=logging.INFO)

class EmailService:
    """Service class for handling email functionality."""
    
    def __init__(self):
        """Initialize the email service."""
        self.smtp_server = None
        self.smtp_port = None
        self.sender_email = None
        self.sender_password = None
        self._load_email_config()

    def _load_email_config(self):
        """Load email configuration from Streamlit secrets."""
        try:
            self.smtp_server = st.secrets["email"]["SMTP_SERVER"]
            self.smtp_port = int(st.secrets["email"]["SMTP_PORT"])
            self.sender_email = st.secrets["email"]["SENDER_EMAIL"]
            self.sender_password = st.secrets["email"]["SENDER_PASSWORD"]
        except Exception as e:
            logging.error(f"Error loading email configuration: {str(e)}")

    def generate_email_pdf(self, content):
        """Generate a professionally formatted PDF for email attachment"""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Add a header with logo placeholder
            pdf.set_font("Arial", "B", 16)
            pdf.cell(190, 10, "Legal Document Analysis", 0, 1, "C")
            pdf.set_font("Arial", "I", 10)
            pdf.cell(190, 5, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, "C")
            pdf.line(10, 25, 200, 25)
            pdf.ln(5)
            
            # Process the content by sections
            content_parts = content.split("\n\n")
            
            for part in content_parts:
                if "DOCUMENT SUMMARY" in part:
                    # Summary section
                    pdf.set_font("Arial", "B", 14)
                    pdf.cell(190, 10, "Document Summary", 0, 1, "L")
                    pdf.set_font("Arial", "", 11)
                    summary_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part
                    pdf.multi_cell(190, 7, summary_content)
                    pdf.ln(5)
                    
                elif "RISK SCORE" in part:
                    # Risk score section
                    pdf.set_font("Arial", "B", 14)
                    pdf.cell(190, 10, "Risk Assessment Score", 0, 1, "L")
                    pdf.set_font("Arial", "", 11)
                    score_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part
                    pdf.multi_cell(190, 7, score_content)
                    pdf.ln(5)
                    
                elif "RISK ANALYSIS" in part:
                    # Risk analysis section
                    pdf.set_font("Arial", "B", 14)
                    pdf.cell(190, 10, "Detailed Risk Analysis", 0, 1, "L")
                    pdf.set_font("Arial", "", 11)
                    analysis_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part
                    pdf.multi_cell(190, 7, analysis_content)
                    pdf.ln(5)
                
                else:
                    # Other content
                    pdf.set_font("Arial", "", 11)
                    pdf.multi_cell(190, 7, part)
                    pdf.ln(3)
            
            # Add footer
            pdf.set_y(-15)
            pdf.set_font("Arial", "I", 8)
            pdf.cell(0, 10, f"Page {pdf.page_no()}/{{nb}}", 0, 0, "C")
            
            pdf_bytes = pdf.output(dest="S").encode("latin1")
            buf = io.BytesIO(pdf_bytes)
            buf.seek(0)
            return buf
            
        except Exception as e:
            logging.error(f"Error generating PDF: {str(e)}")
            return None

    def generate_email_docx(self, content):
        """Generate a professionally formatted Word document for email attachment"""
        try:
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = "Legal Document Analysis"
            doc.core_properties.author = "AI Legal Document Assistant"
            
            # Add header with title
            header = doc.add_heading("Legal Document Analysis Report", level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_run.italic = True
            
            # Add horizontal line
            doc.add_paragraph("_" * 60)
            
            # Process content by sections
            lines = content.split("\n")
            current_section = None
            
            for line in lines:
                if any(section in line for section in [
                    "DOCUMENT SUMMARY", "RISK SCORE", "RISK ANALYSIS", 
                    "COMPLIANCE ANALYSIS", "DETAILED COMPLIANCE REPORT",
                    "KEY FINDINGS", "RECOMMENDATIONS", 
                    "CHAT INTERACTION HISTORY", "DOCUMENT COMPARISON ANALYSIS"
                ]):
                    doc.add_paragraph()
                    heading_text = line.strip().title()
                    heading = doc.add_heading(heading_text, level=2)
                    current_section = line.strip()
                elif line.strip() and not line.startswith("="):
                    p = doc.add_paragraph()
                    
                    if current_section == "RISK ANALYSIS":
                        if line.lower().startswith(("high", "medium", "low")) and ":" in line:
                            priority, description = line.split(":", 1)
                            priority = priority.strip()
                            
                            priority_run = p.add_run(f"{priority}: ")
                            priority_run.bold = True
                            
                            if "high" in priority.lower():
                                priority_run.font.color.rgb = RGBColor(255, 0, 0)
                            elif "medium" in priority.lower():
                                priority_run.font.color.rgb = RGBColor(255, 165, 0)
                            
                            p.add_run(description.strip())
                        else:
                            p.add_run(line)
                    elif current_section == "DOCUMENT COMPARISON ANALYSIS":
                        if "Similarity Score:" in line:
                            p.add_run(line).bold = True
                        elif any(header in line for header in ["Analysis:", "Key Differences:", "Common Elements:"]):
                            p.add_run(line).bold = True
                        else:
                            p.add_run(line)
                    elif current_section == "CHAT INTERACTION HISTORY":
                        if ":" in line:  # This is a chat message
                            role, message = line.split(":", 1)
                            role_run = p.add_run(f"{role}: ")
                            role_run.bold = True
                            p.add_run(message.strip())
                        else:
                            p.add_run(line)
                    else:
                        p.add_run(line)
            
            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph("End of Report")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logging.error(f"Error generating DOCX: {str(e)}")
            return None

    def send_email(self, recipient_email, subject, body, attachment=None, attachment_name=None, attachment_type=None):
        """Send an email with optional attachment."""
        try:
            if not all([self.smtp_server, self.smtp_port, self.sender_email, self.sender_password]):
                return False, "Email configuration is incomplete. Please check your Streamlit secrets."
                
            # Create message
            msg = MIMEMultipart()
            msg['From'] = self.sender_email
            msg['To'] = recipient_email
            msg['Subject'] = subject
            
            # Add body
            msg.attach(MIMEText(body, 'plain'))
            
            # Add attachment if provided
            if attachment and attachment_name and attachment_type:
                attachment_data = attachment
                
                if attachment_type == "pdf":
                    part = MIMEApplication(attachment_data.read() if hasattr(attachment_data, 'read') else attachment_data, Name=attachment_name)
                    part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
                    msg.attach(part)
                elif attachment_type == "docx":
                    part = MIMEApplication(attachment_data.getvalue(), Name=attachment_name)
                    part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
                    msg.attach(part)
                elif attachment_type == "txt":
                    part = MIMEText(attachment_data)
                    part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
                    msg.attach(part)
                    
            # Connect to server and send
            with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                server.starttls()
                server.login(self.sender_email, self.sender_password)
                server.send_message(msg)
                
            return True, "Email sent successfully!"
            
        except Exception as e:
            logging.error(f"Error sending email: {str(e)}")
            return False, f"Error sending email: {str(e)}"

    @staticmethod
    def validate_email(email):
        """Simple email validation"""
        if not email:
            return False
        if "@" not in email or "." not in email:
            return False
        return True

    def prepare_selected_content(self, selected_components: list) -> str:
        """Prepare email content based on selected components."""
        content_parts = []
        
        # Add header
        content_parts.extend([
            "LEGAL DOCUMENT ANALYSIS REPORT",
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "=" * 80,
            "\n"
        ])
        
        # Add selected components
        if "Document Summary" in selected_components and st.session_state.get("summary"):
            content_parts.extend([
                "DOCUMENT SUMMARY",
                "-" * 50,
                st.session_state.summary,
                "\n"
            ])
        
        if "Risk Score" in selected_components and st.session_state.get("risk_score"):
            content_parts.extend([
                "RISK SCORE ANALYSIS",
                "-" * 50,
                f"Overall Risk Score: {st.session_state.risk_score}/100",
                f"Risk Level: {st.session_state.risk_level}",
                "\n"
            ])
        
        if "Risk Analysis" in selected_components and st.session_state.get("risks"):
            content_parts.extend([
                "DETAILED RISK ANALYSIS",
                "-" * 50,
                st.session_state.risks,
                "\n"
            ])
        
        if "Compliance Analysis" in selected_components and st.session_state.get("compliance_results"):
            content_parts.extend([
                "COMPLIANCE ANALYSIS",
                "-" * 50,
                st.session_state.compliance_results,
                "\n"
            ])
            
        if "Detailed Compliance Report" in selected_components and st.session_state.get("compliance_details"):
            content_parts.extend([
                "DETAILED COMPLIANCE REPORT",
                "-" * 50,
                st.session_state.compliance_details,
                "\n"
            ])
        
        if "Key Findings" in selected_components and st.session_state.get("key_findings"):
            content_parts.extend([
                "KEY FINDINGS",
                "-" * 50,
                st.session_state.key_findings,
                "\n"
            ])
        
        if "Recommendations" in selected_components and st.session_state.get("recommendations"):
            content_parts.extend([
                "RECOMMENDATIONS",
                "-" * 50,
                st.session_state.recommendations,
                "\n"
            ])
            
        if "Chat History" in selected_components and st.session_state.get("chat_history"):
            content_parts.extend([
                "CHAT INTERACTION HISTORY",
                "-" * 50,
                self._format_chat_history(st.session_state.chat_history),
                "\n"
            ])
            
        if "Document Comparison" in selected_components and st.session_state.get("comparison_results"):
            content_parts.extend([
                "DOCUMENT COMPARISON ANALYSIS",
                "-" * 50,
                self._format_comparison_results(st.session_state.comparison_results),
                "\n"
            ])
        
        # Add footer
        content_parts.extend([
            "-" * 80,
            "End of Report",
            "\n"
        ])
        
        return "\n".join(content_parts)

    def _format_chat_history(self, chat_history: list) -> str:
        """Format chat history for email content."""
        formatted_chat = []
        for msg in chat_history:
            if isinstance(msg, dict):
                role = msg.get("role", "")
                content = msg.get("content", "")
                formatted_chat.append(f"{role.upper()}: {content}\n")
            elif isinstance(msg, (list, tuple)) and len(msg) >= 2:
                formatted_chat.append(f"{msg[0].upper()}: {msg[1]}\n")
            else:
                formatted_chat.append(str(msg) + "\n")
        return "\n".join(formatted_chat)

    def _format_comparison_results(self, comparison_results: dict) -> str:
        """Format comparison results for email content."""
        formatted_comparison = []
        
        if isinstance(comparison_results, dict):
            # Add similarity score if available
            if "similarity" in comparison_results:
                formatted_comparison.append(f"Similarity Score: {comparison_results['similarity']}%\n")
            
            # Add main analysis
            if "analysis" in comparison_results:
                formatted_comparison.append("Analysis:")
                formatted_comparison.append(comparison_results["analysis"])
                formatted_comparison.append("")
            
            # Add key differences
            if "key_differences" in comparison_results:
                formatted_comparison.append("Key Differences:")
                for category, differences in comparison_results["key_differences"].items():
                    formatted_comparison.append(f"\n{category}:")
                    for diff in differences:
                        if isinstance(diff, dict):
                            if "doc1" in diff and "doc2" in diff:
                                formatted_comparison.append("Document 1: " + diff["doc1"])
                                formatted_comparison.append("Document 2: " + diff["doc2"])
                                formatted_comparison.append("-" * 40)
                            elif "content" in diff:
                                formatted_comparison.append(diff["content"])
                        else:
                            formatted_comparison.append(str(diff))
            
            # Add common elements
            if "common_elements" in comparison_results:
                formatted_comparison.append("\nCommon Elements:")
                for category, elements in comparison_results["common_elements"].items():
                    formatted_comparison.append(f"\n{category}:")
                    for element in elements:
                        if isinstance(element, dict):
                            if "title" in element:
                                formatted_comparison.append(f"- {element['title']}")
                            if "content" in element:
                                formatted_comparison.append(f"  {element['content']}")
                        else:
                            formatted_comparison.append(f"- {element}")
        
        return "\n".join(formatted_comparison)

def email_ui_section():
    """Display email UI section with comprehensive options."""
    st.markdown("### 📧 Email Analysis")
    
    if not st.session_state.get("extracted_text"):
        st.info("Upload and analyze a document to enable email options.")
        return
        
    # Email input
    recipient_email = st.text_input("Recipient Email Address")
    
    if recipient_email:
        email_service = EmailService()
        
        if not email_service.validate_email(recipient_email):
            st.error("Please enter a valid email address.")
            return
            
        # Component selection
        st.markdown("#### Select Components to Include")
        
        available_components = []
        if st.session_state.get("summary"):
            available_components.append("Document Summary")
        if st.session_state.get("risks"):
            available_components.append("Risk Analysis")
        if st.session_state.get("risk_score"):
            available_components.append("Risk Score")
        if st.session_state.get("compliance_results"):
            available_components.append("Compliance Analysis")
        if st.session_state.get("compliance_details"):
            available_components.append("Detailed Compliance Report")
        if st.session_state.get("key_findings"):
            available_components.append("Key Findings")
        if st.session_state.get("recommendations"):
            available_components.append("Recommendations")
        if st.session_state.get("chat_history"):
            available_components.append("Chat History")
        if st.session_state.get("comparison_results"):
            available_components.append("Document Comparison")
            
        if not available_components:
            st.info("Analyze the document first to enable email options.")
            return
            
        selected_components = st.multiselect(
            "Choose sections to include in the email:",
            options=available_components,
            default=available_components,
            help="Select which analysis components you want to include in the email"
        )
        
        if not selected_components:
            st.warning("Please select at least one component to include.")
            return
            
        # Email customization
        st.markdown("#### Email Options")
        
        subject = st.text_input(
            "Email Subject",
            value="Legal Document Analysis Report",
            help="Customize the email subject line"
        )
        
        # Prepare email content based on selection
        email_content = email_service.prepare_selected_content(selected_components)
        
        if email_content:
            # Format selection
            attachment_format = st.selectbox(
                "Attachment Format",
                ["PDF", "DOCX", "TXT"],
                help="Choose the format for the analysis attachment"
            )
            
            if st.button("Send Email"):
                try:
                    with st.spinner("Sending email..."):
                        # Generate attachment
                        if attachment_format == "PDF":
                            attachment = email_service.generate_email_pdf(email_content)
                            attachment_name = f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                            mime_type = "application/pdf"
                        elif attachment_format == "DOCX":
                            attachment = email_service.generate_email_docx(email_content)
                            attachment_name = f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        else:  # TXT
                            attachment = email_content.encode('utf-8')
                            attachment_name = f"legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.txt"
                            mime_type = "text/plain"
                        
                        # Send email
                        email_service.send_email(
                            recipient_email=recipient_email,
                            subject=subject,
                            body="Please find attached the legal document analysis report.",
                            attachment=attachment,
                            attachment_name=attachment_name,
                            attachment_type=mime_type
                        )
                        
                        st.success("Email sent successfully! 📨")
                except Exception as e:
                    st.error(f"Error sending email: {str(e)}")
                    logging.error(f"Email error: {str(e)}", exc_info=True) 