# In the upload section
# In the upload section
# st.button("Analyze Document", key="analyze_doc_button")
"""AI Legal Document Assistant

A Streamlit application for analyzing, summarizing, and assessing risks in legal documents.
This application uses Google's Gemini API to provide AI-powered document analysis.

Features:
- Document upload and text extraction (PDF, DOCX, TXT)
- AI-powered document summarization
- Legal risk identification and assessment
- Interactive chat with documents
- Risk visualization with charts and metrics
- Export options (PDF, DOCX, TXT) for analysis results

Deployment Requirements:
- Python 3.9+
- Streamlit Cloud account
- Google Gemini API key (configured in .streamlit/secrets.toml)
- Required packages in requirements.txt

Created for Legal professionals and businesses to streamline document analysis.
"""

import streamlit as st
import logging
import smtplib
import io
import os
import json
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pdf_processor import extract_text_from_pdf
from ai_analyzer import identify_risks, summarize_document, chat_with_document
from utils_state_management import initialize_session_state, reset_analysis_state, update_document_state
from fpdf import FPDF
from docx import Document
from datetime import datetime
from compliance_analyzer import ComplianceAnalyzer
# Import GDPR compliance features
from gdpr_compliance import (
    GDPRCompliance,
    show_gdpr_consent_banner,
    add_privacy_policy_footer,
    show_gdpr_info_iframe,
    show_privacy_policy
)
# Import email functionality
from email_service import EmailService, email_ui_section
try:
    from email_service import email_ui_section
    email_enabled = True
except ImportError:
    email_enabled = False
    logging.warning("Email service module not found. Email functionality disabled.")
# Import document comparison module
import document_comparison as doc_comp
from utils_file_processor import extract_text_from_uploaded_file
from document_analyzer import DocumentAnalyzer
from risk_analyzer import RiskAnalyzer
from compliance_checker import ComplianceChecker
from document_comparer import DocumentComparer
from chat_handler import ChatHandler
from export_handler import ExportHandler

st.set_page_config(
    page_title="AI Legal Document Assistant",
    page_icon="⚖️",
    layout="wide"
)

# Initialize logging
logging.basicConfig(level=logging.INFO)

# Initialize session state variables
initialize_session_state()

# Initialize components
document_analyzer = DocumentAnalyzer()
risk_analyzer = RiskAnalyzer()
compliance_checker = ComplianceChecker()
document_comparer = DocumentComparer()
chat_handler = ChatHandler()
export_handler = ExportHandler()
gdpr_compliance = GDPRCompliance()

# Override any None values to empty strings for better handling
if st.session_state.extracted_text is None:
    st.session_state.extracted_text = ""
if st.session_state.summary is None:
    st.session_state.summary = ""
if st.session_state.risks is None:
    st.session_state.risks = ""

# Set current year for privacy policy footer
st.session_state['current_year'] = datetime.now().year

# Check if we should show the privacy policy
if st.session_state.get('show_privacy_policy', False):
    show_privacy_policy()
    st.stop()

# Show GDPR consent banner if consent not given
show_gdpr_consent_banner()

# Initialize compliance analyzer - with improved error handling
try:
    compliance_rules_path = os.path.join(os.path.dirname(__file__), "compliance", "compliance_rules.json")
    if os.path.exists(compliance_rules_path):
        compliance_analyzer = ComplianceAnalyzer(config_path=compliance_rules_path)
    else:
        logging.warning(f"Compliance rules file not found at {compliance_rules_path}, using default rules.")
        compliance_analyzer = ComplianceAnalyzer()
except Exception as e:
    logging.error(f"Error initializing compliance analyzer: {str(e)}")
    compliance_analyzer = ComplianceAnalyzer()  # Fallback to default rules

def extract_text_from_uploaded_file(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(uploaded_file)
    elif ext in [".doc", ".docx"]:
        try:
            document = Document(uploaded_file)
            fullText = [para.text for para in document.paragraphs]
            return "\n".join(fullText)
        except Exception as e:
            logging.error(f"Error reading DOC/DOCX: {str(e)}")
            return None
    elif ext == ".txt":
        try:
            return uploaded_file.getvalue().decode("utf-8")
        except Exception as e:
            logging.error(f"Error reading TXT: {str(e)}")
            return None
    else:
        return None

def generate_txt(content):
    """Generate a well-formatted text document for download"""
    # Format the content with clear section dividers
    lines = content.split("\n")
    formatted_lines = []
    
    # Add a header
    formatted_lines.append("=" * 80)
    formatted_lines.append("LEGAL DOCUMENT ANALYSIS REPORT")
    formatted_lines.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    formatted_lines.append("=" * 80)
    formatted_lines.append("")
    
    # Process content with proper spacing
    in_section = False
    section_title = ""
    
    for line in lines:
        # Check for section headers
        if "DOCUMENT SUMMARY" in line or "RISK ANALYSIS" in line or "RISK SCORE" in line:
            # Add spacing between sections
            if in_section:
                formatted_lines.append("")
            
            in_section = True
            section_title = line.strip()
            
            # Add section header with emphasis
            formatted_lines.append("")
            formatted_lines.append(section_title)
            formatted_lines.append("-" * len(section_title))
            
        elif line.strip() == "=" * 50:
            # Skip the original separator lines
            continue
        else:
            # Regular content - ensure proper line spacing
            if line.strip():
                formatted_lines.append(line)
    
    # Add footer
    formatted_lines.append("")
    formatted_lines.append("-" * 80)
    formatted_lines.append("End of Report")
    
    return "\n".join(formatted_lines)

def generate_docx(content):
    """Generate a professionally formatted Word document for download"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Legal Document Analysis"
    doc.core_properties.author = "AI Legal Document Assistant"
    
    # Add header with title
    header = doc.add_heading("Legal Document Analysis Report", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.italic = True
    
    # Add horizontal line
    doc.add_paragraph("_" * 60)
    
    # Process content by sections
    lines = content.split("\n")
    current_section = None
    
    for line in lines:
        if "DOCUMENT SUMMARY" in line:
            # Add Summary section with formatting
            doc.add_paragraph()  # Add some spacing
            heading = doc.add_heading("Document Summary", level=2)
            current_section = "summary"
            
        elif "RISK SCORE" in line:
            # Add Risk Score section with formatting
            doc.add_paragraph()  # Add some spacing
            heading = doc.add_heading("Risk Assessment Score", level=2)
            current_section = "risk_score"
            
        elif "RISK ANALYSIS" in line:
            # Add Risk Analysis section with formatting
            doc.add_paragraph()  # Add some spacing
            heading = doc.add_heading("Detailed Risk Analysis", level=2)
            current_section = "risk_analysis"
            
        elif line.strip() and not line.startswith("=") and not line.startswith("-"):
            # Regular content - check for potential risk items
            p = doc.add_paragraph()
            
            # Check if this might be a risk item with priority
            if current_section == "risk_analysis":
                if line.lower().startswith(("high", "medium", "low")) and ":" in line:
                    # This is a risk item with priority
                    priority, description = line.split(":", 1)
                    priority = priority.strip()
                    
                    # Format based on priority
                    priority_run = p.add_run(f"{priority}: ")
                    priority_run.bold = True
                    
                    if "high" in priority.lower():
                        priority_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for high
                    elif "medium" in priority.lower():
                        priority_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange for medium
                    
                    p.add_run(description.strip())
                else:
                    p.add_run(line)
            else:
                p.add_run(line)
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph("End of Report")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document to a BytesIO object
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_pdf(content):
    """Generate a professionally formatted PDF for download"""
    from fpdf import FPDF
    from datetime import datetime
    
    class PDF(FPDF):
        def header(self):
            # Add logo (placeholder)
            # self.image('logo.png', 10, 8, 33)
            
            # Set font for header
            self.set_font('Arial', 'B', 16)
            
            # Title
            self.cell(0, 10, 'Legal Document Analysis Report', 0, 1, 'C')
            
            # Date
            self.set_font('Arial', 'I', 10)
            self.cell(0, 5, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, 'C')
            
            # Line break
            self.ln(5)
            
            # Draw a line
            self.line(10, 25, 200, 25)
            
            # Line break after header
            self.ln(10)
            
        def footer(self):
            # Position at 1.5 cm from bottom
            self.set_y(-15)
            
            # Set font for footer
            self.set_font('Arial', 'I', 8)
            
            # Page number
            self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')
    
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Process content by sections
    content_parts = content.split("\n\n")
    current_section = None
    
    for part in content_parts:
        if "DOCUMENT SUMMARY" in part:
            # Summary section
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Document Summary", 0, 1, "L")
            pdf.set_font("Arial", "", 11)
            
            # Get content after header
            summary_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("DOCUMENT SUMMARY", "").strip()
            
            # Add the summary text
            pdf.multi_cell(0, 7, sanitize_text_for_pdf(summary_content))
            pdf.ln(5)
            
        elif "RISK SCORE" in part:
            # Risk score section
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Risk Assessment Score", 0, 1, "L")
            pdf.set_font("Arial", "", 11)
            
            # Get content after header
            score_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("RISK SCORE", "").strip()
            
            # Parse score lines
            for line in score_content.split("\n"):
                if "Score:" in line or "Risk Level:" in line:
                    pdf.set_font("Arial", "B", 11)
                    pdf.cell(0, 7, line, 0, 1)
                    pdf.set_font("Arial", "", 11)
                else:
                    pdf.cell(0, 7, line, 0, 1)
                    
            pdf.ln(5)
            
        elif "RISK ANALYSIS" in part:
            # Risk analysis section
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Detailed Risk Analysis", 0, 1, "L")
            pdf.set_font("Arial", "", 11)
            
            # Get content after header
            analysis_content = part.split("-" * 30)[1].strip() if "-" * 30 in part else part.replace("RISK ANALYSIS", "").strip()
            
            # Check for priority sections
            if "HIGH PRIORITY RISKS" in analysis_content:
                # Handle structured risk content
                sections = analysis_content.split("\n\n")
                for section in sections:
                    if "HIGH PRIORITY RISKS" in section:
                        pdf.set_font("Arial", "B", 12)
                        pdf.set_text_color(255, 0, 0)  # Red
                        pdf.cell(0, 10, "High Priority Risks", 0, 1)
                        pdf.set_text_color(0, 0, 0)  # Black
                        pdf.set_font("Arial", "", 11)
                        
                        # Add each risk
                        for line in section.split("\n"):
                            if not line.strip() or "HIGH PRIORITY RISKS" in line:
                                continue
                            pdf.multi_cell(0, 7, sanitize_text_for_pdf(line))
                            
                    elif "MEDIUM PRIORITY RISKS" in section:
                        pdf.set_font("Arial", "B", 12)
                        pdf.set_text_color(255, 165, 0)  # Orange
                        pdf.cell(0, 10, "Medium Priority Risks", 0, 1)
                        pdf.set_text_color(0, 0, 0)  # Black
                        pdf.set_font("Arial", "", 11)
                        
                        # Add each risk
                        for line in section.split("\n"):
                            if not line.strip() or "MEDIUM PRIORITY RISKS" in line:
                                continue
                            pdf.multi_cell(0, 7, sanitize_text_for_pdf(line))
                            
                    elif "LOW PRIORITY RISKS" in section:
                        pdf.set_font("Arial", "B", 12)
                        pdf.set_text_color(0, 128, 0)  # Green
                        pdf.cell(0, 10, "Low Priority Risks", 0, 1)
                        pdf.set_text_color(0, 0, 0)  # Black
                        pdf.set_font("Arial", "", 11)
                        
                        # Add each risk
                        for line in section.split("\n"):
                            if not line.strip() or "LOW PRIORITY RISKS" in line:
                                continue
                            pdf.multi_cell(0, 7, sanitize_text_for_pdf(line))
            else:
                # Unstructured content
                pdf.multi_cell(0, 7, sanitize_text_for_pdf(analysis_content))
                
            pdf.ln(5)
            
        else:
            # Other content
            pdf.set_font("Arial", "", 11)
            pdf.multi_cell(0, 7, sanitize_text_for_pdf(part))
            pdf.ln(3)
    
    # Get the PDF as bytes
    pdf_str = pdf.output(dest="S")
    if isinstance(pdf_str, str):
        pdf_bytes = pdf_str.encode("latin1")
    else:
        pdf_bytes = pdf_str
        
    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)
    return buf

def sanitize_text_for_pdf(text):
    """Clean and sanitize text for PDF output to avoid encoding issues"""
    if not text:
        return ""
    
    # Replace problematic Unicode characters
    replacements = {
        '\u2013': '-',  # en dash
        '\u2014': '--', # em dash
        '\u2018': "'",  # left single quote
        '\u2019': "'",  # right single quote
        '\u201c': '"',  # left double quote
        '\u201d': '"',  # right double quote
        '\u2022': '*',  # bullet
        '\u2026': '...', # ellipsis
        '\u00a0': ' ',  # non-breaking space
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    # Replace any other non-Latin1 characters with their closest ASCII equivalent or '?'
    clean_text = ""
    for char in text:
        if ord(char) < 256:
            clean_text += char
        else:
            clean_text += '?'
            
    return clean_text

def calculate_risk_score(risks_text):
    """Calculate a simplified risk score to avoid connection errors"""
    if not risks_text or not isinstance(risks_text, str):
        return 0, 0, 0, 0
    
    # Use simple keyword counting - limit processing to avoid timeouts
    text_lower = risks_text.lower()
    max_text_length = min(len(text_lower), 10000)  # Cap text length for processing
    text_sample = text_lower[:max_text_length]
    
    # Count simple keyword occurrences
    high_count = sum(text_sample.count(word) for word in ['critical', 'severe', 'high risk', 'significant', 'major', 'serious'])
    medium_count = sum(text_sample.count(word) for word in ['moderate', 'medium', 'potential', 'possible', 'concerning'])
    low_count = len(text_sample.split('.')) - high_count - medium_count
    low_count = max(0, low_count)
    
    total = high_count + medium_count + low_count
    if total == 0:
        return 0, 0, 0, 0
    
    # Simple calculation
    score = min(100, round((high_count * 3 + medium_count * 2 + low_count) / max(total, 1) * 20))
    
    return score, high_count, medium_count, low_count

def display_risk_score(risks_text):
    """Display ultra-simplified risk score to prevent connection errors"""
    if not risks_text or not isinstance(risks_text, str):
        st.info("No risk data available for scoring.")
        return
    
    try:
        # Use the simplified calculation
        score, high_count, medium_count, low_count = calculate_risk_score(risks_text)
        
        # Simple text display without complex components
        st.subheader("Risk Score Analysis")
        st.write(f"Overall Risk Score: {score}/100")
        
        # Basic risk level display
        if score >= 70:
            st.error("🔴 High Risk Level")
        elif score >= 40:
            st.warning("🟠 Medium Risk Level")
        else:
            st.success("🟢 Low Risk Level")
        
        # Simple text for counts instead of metrics
        st.write("Risk Breakdown:")
        st.write(f"- High Priority Risks: {high_count}")
        st.write(f"- Medium Priority Risks: {medium_count}")
        st.write(f"- Low Priority Risks: {low_count}")
    
    except Exception as e:
        # Fallback if any error occurs
        st.error(f"Error calculating risk score: {str(e)}")
        st.info("Please try again with a different document.")

def display_risks_with_filters(risks_text):
    """Display risks with filtering options from text"""
    if not risks_text or not isinstance(risks_text, str):
        st.warning("No risks identified.")
        return
    
    # Split the text into sentences
    sentences = [s.strip() for s in risks_text.split('.') if s.strip()]
    
    if not sentences:
        st.warning("No clear risk statements identified.")
        return
        
    # Count risks
    st.write(f"Total Potential Risks Found: {len(sentences)}")
    
    # Add filter options
    priority_filter = st.multiselect(
        "Filter by Priority",
        ["High", "Medium", "Low"],
        default=["High", "Medium", "Low"]
    )
    
    # Categorize sentences by risk level
    high_risk_keywords = ['critical', 'severe', 'high risk', 'significant', 'major', 'serious']
    medium_risk_keywords = ['moderate', 'medium', 'potential', 'possible', 'concerning']
    
    high_risks = []
    medium_risks = []
    low_risks = []
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        if any(keyword in sentence_lower for keyword in high_risk_keywords):
            high_risks.append(sentence)
        elif any(keyword in sentence_lower for keyword in medium_risk_keywords):
            medium_risks.append(sentence)
        elif len(sentence.split()) > 5:  # Only include as low risk if it's a substantial sentence
            low_risks.append(sentence)
    
    # Display filtered risks
    if "High" in priority_filter and high_risks:
        st.markdown("#### 🔴 High Priority Risks")
        for idx, risk in enumerate(high_risks, 1):
            st.error(f"{idx}. {risk}")
    
    if "Medium" in priority_filter and medium_risks:
        st.markdown("#### 🟠 Medium Priority Risks")
        for idx, risk in enumerate(medium_risks, 1):
            st.warning(f"{idx}. {risk}")
    
    if "Low" in priority_filter and low_risks:
        st.markdown("#### 🟢 Low Priority Risks")
        for idx, risk in enumerate(low_risks, 1):
            st.info(f"{idx}. {risk}")

def display_risk_summary(risks_text):
    """Display a summarized view of the risk assessment"""
    if not risks_text or not isinstance(risks_text, str):
        st.warning("No risk data available for summary.")
        return
    
    # Calculate score
    score, high_count, medium_count, low_count = calculate_risk_score(risks_text)
    
    # Create a card-like summary
    st.subheader("Risk Assessment Summary")
    
    # Create two columns
    col1, col2 = st.columns(2)
    
    with col1:
        # Display metrics
        st.metric("Overall Risk Score", f"{score}/100")
        
        # Risk level with color
        if score >= 70:
            st.markdown("### 🔴 High Risk Level")
        elif score >= 40:
            st.markdown("### 🟠 Medium Risk Level")
        else:
            st.markdown("### 🟢 Low Risk Level")
    
    with col2:
        # Risk counts
        st.write("Risk Breakdown:")
        st.error(f"High Priority Risks: {high_count}")
        st.warning(f"Medium Priority Risks: {medium_count}")
        st.info(f"Low Priority Risks: {low_count}")
    
    # Key recommendations based on risk level
    st.markdown("### Recommendations:")
    if score >= 70:
        st.markdown("- 🚨 Immediate legal review highly recommended")
        st.markdown("- 🚨 Address high priority risks before proceeding")
        st.markdown("- 🚨 Consider professional legal consultation")
    elif score >= 40:
        st.markdown("- ⚠️ Review document carefully before proceeding")
        st.markdown("- ⚠️ Address medium priority risks")
        st.markdown("- ⚠️ Consider additional review for specific sections")
    else:
        st.markdown("- ✅ Document appears to have low risk")
        st.markdown("- ✅ Standard review procedures recommended")
        st.markdown("- ✅ Monitor for changes that might increase risk")

def display_risk_visualizer(risks_text):
    """Display a simple visualization of the risks to avoid connection errors"""
    if not risks_text or not isinstance(risks_text, str):
        st.info("No risk data available to visualize.")
        return
    
    try:
        # Use the simplified calculation
        score, high_count, medium_count, low_count = calculate_risk_score(risks_text)
        
        st.subheader("Risk Visualization")
        
        # Create a simple progress bar for overall risk
        st.write("Overall Risk Level:")
        if score >= 70:
            color = "red"
            level = "High Risk 🔴"
        elif score >= 40:
            color = "orange"
            level = "Medium Risk 🟠"
        else:
            color = "green"
            level = "Low Risk 🟢"
            
        st.markdown(f"### {level}")
        st.progress(min(score/100, 1.0))
        
        # Risk distribution
        st.write("Risk Distribution:")
        
        total = high_count + medium_count + low_count
        if total > 0:
            # Calculate percentages
            high_pct = int((high_count / total) * 100)
            medium_pct = int((medium_count / total) * 100)
            low_pct = 100 - high_pct - medium_pct
            
            # Create simplified horizontal bar chart using markdown
            st.markdown("**High Priority** 🔴")
            st.progress(high_count / max(total, 1))
            st.markdown(f"{high_count} issues ({high_pct}%)")
            
            st.markdown("**Medium Priority** 🟠")
            st.progress(medium_count / max(total, 1))
            st.markdown(f"{medium_count} issues ({medium_pct}%)")
            
            st.markdown("**Low Priority** 🟢")
            st.progress(low_count / max(total, 1))
            st.markdown(f"{low_count} issues ({low_pct}%)")
            
            # Top risk areas (simplified)
            st.subheader("Key Risk Indicators")
            
            # Find most common risk keywords
            risk_keywords = [
                ("non-compliance", "regulatory requirements"),
                ("liability", "legal exposure"),
                ("breach", "contract terms"),
                ("confidentiality", "disclosure"),
                ("warranty", "guarantees"),
                ("termination", "cancellation"),
                ("intellectual property", "IP rights"),
                ("dispute", "resolution"),
                ("payment", "financial terms")
            ]
            
            # Count keyword occurrences
            keyword_counts = []
            for kw_pair in risk_keywords:
                count = sum(risks_text.lower().count(kw) for kw in kw_pair)
                if count > 0:
                    keyword_counts.append((f"{kw_pair[0]}/{kw_pair[1]}", count))
            
            # Display top 5 risk areas
            if keyword_counts:
                keyword_counts.sort(key=lambda x: x[1], reverse=True)
                for kw, count in keyword_counts[:5]:
                    st.write(f"• {kw}: {count} mentions")
            else:
                st.info("No specific risk areas identified.")
                
    except Exception as e:
        st.error(f"Error generating visualization: {str(e)}")
        st.info("Please try again with a different document.")

def analyze_document():
    st.subheader("Document Analysis")
    if st.session_state.get("extracted_text"):
        if st.button("Generate Summary", key="generate_summary_button"):
            with st.spinner("Analyzing document..."):
                summary = summarize_document(st.session_state.extracted_text)
                st.session_state.summary = summary if summary else "No summary generated."
        
        if st.session_state.get("summary"):
            st.markdown("### 📝 Document Summary")
            st.write(st.session_state.summary)
            
    else:
        st.info("Please upload a document in the Upload tab first.")

def compare_documents():
    st.subheader("Compare Documents")
    
    # First document is automatically selected from the main document
    if 'uploaded_docs' not in st.session_state or not st.session_state.uploaded_docs:
        st.warning("Please upload at least one document in the Upload tab first.")
        return
    
    # Get the main document (first document or user-selected main document)
    if 'main_doc_id' in st.session_state and st.session_state.main_doc_id in st.session_state.uploaded_docs:
        main_doc_id = st.session_state.main_doc_id
    else:
        # Default to the first document if no main document is set
        main_doc_id = list(st.session_state.uploaded_docs.keys())[0]
    
    # Display the main document info
    st.info(f"**First Document:** {st.session_state.uploaded_docs[main_doc_id]['name']} (from Upload tab)")
    
    # Option for selecting second document: either from existing uploads or new upload
    st.subheader("Select Second Document")
    compare_option = st.radio(
        "Choose second document source:",
        ["Upload a new document", "Use an existing uploaded document"]
    )
    
    doc2 = None
    doc2_text = None
    doc2_name = None
    
    if compare_option == "Upload a new document":
        uploaded_comparison_file = st.file_uploader(
            "Upload a document to compare with the main document", 
            type=["pdf", "txt"],
            key="comparison_uploader"
        )
        
        if uploaded_comparison_file:
            # Extract text from the newly uploaded document
            doc2_text = extract_text_from_uploaded_file(uploaded_comparison_file)
            doc2_name = uploaded_comparison_file.name
            
            # Show confirmation
            st.success(f"Uploaded: {doc2_name}")
            
    else:  # Use existing document
        # Filter out the main document from the selection
        remaining_docs = [doc for doc in st.session_state.uploaded_docs.keys() if doc != main_doc_id]
        
        if remaining_docs:
            doc2 = st.selectbox(
                "Select document to compare", 
                options=remaining_docs,
                format_func=lambda x: st.session_state.uploaded_docs[x]['name']
            )
            
            if doc2:
                doc2_text = st.session_state.uploaded_docs[doc2]['text']
                doc2_name = st.session_state.uploaded_docs[doc2]['name']
                
        else:
            st.warning("No other documents available. Please upload another document or use the direct upload option.")
    
    # Comparison type selection
    comparison_type = st.selectbox(
        "Comparison Focus",
        options=["General Comparison", "Legal Clauses", "Compliance Elements", "Risk Factors"],
        help="Select what aspect of the documents you want to focus on comparing"
    )
    
    # Add a note about the summarization content
    st.info("The comparison will automatically include the document summary to highlight key differences.")
    
    # Only enable the compare button if we have two documents
    compare_button_enabled = bool(doc2_text or doc2)
    
    if st.button("Compare Documents", disabled=not compare_button_enabled):
        if not doc2_text and not doc2:
            st.error("Please select or upload a second document to compare.")
            return
            
        with st.spinner("Comparing documents..."):
            # Get the main document text
            main_doc_text = st.session_state.uploaded_docs[main_doc_id]['text']
            
            # Also get the document summary if available
            main_doc_summary = ""
            if 'summaries' in st.session_state and main_doc_id in st.session_state.summaries:
                main_doc_summary = st.session_state.summaries[main_doc_id]
                
            # Combine main document text with its summary for better comparison
            enhanced_main_doc_text = f"DOCUMENT SUMMARY:\n{main_doc_summary}\n\nFULL TEXT:\n{main_doc_text}"
            
            # Prepare documents for comparison
            doc_texts = {
                'main_doc': enhanced_main_doc_text,
                'second_doc': doc2_text
            }
            
            doc_names = {
                'main_doc': st.session_state.uploaded_docs[main_doc_id]['name'],
                'second_doc': doc2_name
            }
            
            # Call the comparison function from our module
            results = doc_comp.compare_documents(doc_texts, doc_names, comparison_type)
            
            if "error" in results:
                st.error(results["error"])
            else:
                # Display the comparison results
                st.subheader("Comparison Results")
                
                # Display the similarity percentage
                if "similarity" in results:
                    st.metric("Document Similarity", f"{results['similarity']}%")
                
                # Display the main analysis
                st.markdown("### Analysis")
                st.write(results["analysis"])
                
                # Display key differences in an expandable section
                with st.expander("Key Differences", expanded=True):
                    if results["key_differences"]:
                        for category, differences in results["key_differences"].items():
                            st.markdown(f"#### {category}")
                            
                            for diff in differences:
                                if "doc1" in diff and "doc2" in diff:
                                    cols = st.columns(2)
                                    cols[0].markdown(f"**{doc_names['main_doc']}**")
                                    cols[0].write(diff["doc1"])
                                    
                                    cols[1].markdown(f"**{doc_names['second_doc']}**")
                                    cols[1].write(diff["doc2"])
                                    
                                    st.markdown("---")
                                else:
                                    st.write(diff.get("content", ""))
                    else:
                        st.write("No significant differences identified.")
                
                # Display common elements in an expandable section
                with st.expander("Common Elements", expanded=True):
                    if results["common_elements"]:
                        for category, elements in results["common_elements"].items():
                            st.markdown(f"#### {category}")
                            
                            for element in elements:
                                st.markdown(f"**{element.get('title', 'Element')}**")
                                st.write(element.get("content", ""))
                                st.markdown("---")
                    else:
                        st.write("No common elements identified.")

st.title("📜 AI Legal Document Assistant")
st.markdown("**Upload a legal document** to analyze, summarize, and chat with it. Powered by Gemini AI.")

# Create tabs for different sections
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📂 Document Upload", "📊 Analysis", "⚠️ Risk Assessment", "📋 Compliance", "💬 Chat", "📈 Compare Documents"])

with tab1:
    st.subheader("Upload Document")
    uploaded_file = st.file_uploader("Upload a legal document", type=["pdf", "txt"])
    
    if uploaded_file is not None:
        # Extract text and update session state
        text = extract_text_from_uploaded_file(uploaded_file)
        st.session_state.extracted_text = text
        
        # Also save to uploaded_docs for document comparison feature
        file_id = uploaded_file.name
        st.session_state.uploaded_docs[file_id] = {
            'name': uploaded_file.name,
            'text': text,
            'type': uploaded_file.type
        }
        
        # Set this as the main document if none is set
        if st.session_state.main_doc_id is None:
            st.session_state.main_doc_id = file_id
        
        st.success(f"Document uploaded: {uploaded_file.name}")
        
        with st.expander("View Document Text"):
            st.write(text[:1000] + "..." if len(text) > 1000 else text)
    else:
        if st.session_state.get("extracted_text"):
            # Show previous analysis if document was already uploaded
            st.info("Using previously uploaded document. Upload another document to replace it.")

with tab2:
    document_analyzer.analyze_document_ui()

with tab3:
    risk_analyzer.risk_assessment_ui()

with tab4:
    compliance_checker.compliance_check_ui()

with tab5:
    chat_handler.chat_ui()

with tab6:
    compare_documents()

# Sidebar with Export and Email Options
with st.sidebar:
    export_handler.export_ui()
    
    # Email section
    st.markdown("---")
    
    # Check if email is configured
    email_configured = False
    try:
        if "email" in st.secrets and all(key in st.secrets["email"] for key in ["SMTP_SERVER", "SMTP_PORT", "SENDER_EMAIL", "SENDER_PASSWORD"]):
            email_configured = True
    except Exception as e:
        logging.error(f"Error checking email configuration: {str(e)}")
    
    if email_configured:
        if st.session_state.get("summary") or st.session_state.get("risks"):
            # Email UI is in a separate function in email_service.py
            try:
                from email_service import email_ui_section
                email_ui_section()
            except Exception as e:
                st.error(f"Error with email functionality: {str(e)}")
                logging.error(f"Email error details: {str(e)}", exc_info=True)
        else:
            st.info("Upload a document and analyze it to enable email options.")
    else:
        st.warning("""
        Email configuration not found. Add these to your `.streamlit/secrets.toml` file:
        ```
        [email]
        SMTP_SERVER = "smtp.gmail.com"
        SMTP_PORT = 587
        SENDER_EMAIL = "your-email@gmail.com"
        SENDER_PASSWORD = "your-app-password"
        ```
        """)

# GDPR Information iframe (optional)
if st.session_state.get('show_gdpr_iframe', False):
    show_gdpr_info_iframe()

# Add privacy policy footer
add_privacy_policy_footer()
